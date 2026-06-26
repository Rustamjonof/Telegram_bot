"""
Start Daily Bot - AI Funksiyalari
"""

import os
import json
import aiohttp
import asyncio
from datetime import datetime
from aiogram import types, F
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from database import (
    update_user, check_subscription,
    can_use_ai, log_ai_request, TARIFF_LIMITS,
    has_free_usage, mark_free_usage,
    has_enough_balance, deduct_balance, get_balance
)
import importlib
import sys

def get_user(uid):
    main_module = sys.modules.get('__main__')
    if main_module and hasattr(main_module, 'get_user'):
        return main_module.get_user(uid)
    from database import get_user as db_get_user
    return db_get_user(uid)

ADMIN_ID      = int(os.getenv("ADMIN_ID", "5236920689"))
OPENAI_API    = os.getenv("OPENAI_API", "")        # OpenAI API kaliti (env)
GROQ_API_KEY  = os.getenv("GROQ_API_KEY", "")      # Groq API kaliti (env)
GROQ_MODEL    = "llama3-70b-8192"
GROQ_URL      = "https://api.groq.com/openai/v1/chat/completions"
OPENAI_URL    = "https://api.openai.com/v1/chat/completions"
DALLE_URL     = "https://api.openai.com/v1/images/generations"

# ============================================================
# TARJIMALAR
# ============================================================
AI_MENU_TITLE = {
    "uz": "🤖 AI Yordamchi\n\nQaysi funksiyadan foydalanmoqchisiz?",
    "en": "🤖 AI Assistant\n\nWhich function would you like to use?",
    "ru": "🤖 AI Помощник\n\nКакую функцию хотите использовать?",
    "ar": "🤖 مساعد AI\n\nأي وظيفة تريد استخدامها؟",
    "tr": "🤖 AI Asistan\n\nHangi işlevi kullanmak istersiniz?",
    "es": "🤖 Asistente AI\n\n¿Qué función deseas usar?",
    "hi": "🤖 AI सहायक\n\nआप कौन सा फ़ंक्शन उपयोग करना चाहते हैं?",
}

AI_FUNCTIONS = {
    "uz": {
        "image":      "🖼 Rasm yaratish",
        "present":    "📊 Taqdimot/Slayd",
        "referat":    "📝 Referat",
        "tezis":      "📄 Tezis",
        "article":    "📰 Maqola",
        "mustaqil":   "🗂 Mustaqil ish",
        "kurs":       "🎓 Kurs ishi",
        "keys":       "💼 Keys (Case study)",
        "resume":     "📋 Rezyume",
        "texmap":     "🗺 Texnologik xarita",
        "quiz":       "❓ Quiz/Test",
        "crossword":  "🔠 Krossvord",
        "translate":  "🌐 Tarjima",
        "essay":      "✍️ Insho/Esse",
        "flashcard":  "🃏 Flesh karta",
        "convert":    "🔄 Fayl formati",
        "chat":       "💬 Oddiy chat",
    },
    "en": {
        "image":      "🖼 Image generation",
        "present":    "📊 Presentation/Slides",
        "referat":    "📝 Report/Essay",
        "tezis":      "📄 Thesis",
        "article":    "📰 Article",
        "mustaqil":   "🗂 Independent work",
        "kurs":       "🎓 Course work",
        "keys":       "💼 Case study",
        "resume":     "📋 Resume",
        "texmap":     "🗺 Tech map",
        "quiz":       "❓ Quiz/Test",
        "crossword":  "🔠 Crossword",
        "translate":  "🌐 Translation",
        "essay":      "✍️ Essay",
        "flashcard":  "🃏 Flashcard",
        "convert":    "🔄 File convert",
        "chat":       "💬 Chat",
    },
    "ru": {
        "image":      "🖼 Генерация изображений",
        "present":    "📊 Презентация/Слайды",
        "referat":    "📝 Реферат",
        "tezis":      "📄 Тезис",
        "article":    "📰 Статья",
        "mustaqil":   "🗂 Самостоятельная работа",
        "kurs":       "🎓 Курсовая работа",
        "keys":       "💼 Кейс (Case study)",
        "resume":     "📋 Резюме",
        "texmap":     "🗺 Технологическая карта",
        "quiz":       "❓ Викторина/Тест",
        "crossword":  "🔠 Кроссворд",
        "translate":  "🌐 Перевод",
        "essay":      "✍️ Эссе",
        "flashcard":  "🃏 Флэш-карта",
        "convert":    "🔄 Конвертация файлов",
        "chat":       "💬 Чат",
    },
    "ar": {
        "image":      "🖼 توليد الصور",
        "present":    "📊 عرض تقديمي",
        "referat":    "📝 تقرير",
        "tezis":      "📄 أطروحة",
        "article":    "📰 مقال",
        "mustaqil":   "🗂 عمل مستقل",
        "kurs":       "🎓 عمل دراسي",
        "keys":       "💼 دراسة حالة",
        "resume":     "📋 سيرة ذاتية",
        "texmap":     "🗺 خريطة تقنية",
        "quiz":       "❓ اختبار",
        "crossword":  "🔠 كلمات متقاطعة",
        "translate":  "🌐 ترجمة",
        "essay":      "✍️ مقال إنشائي",
        "flashcard":  "🃏 بطاقة تعليمية",
        "convert":    "🔄 تحويل الملفات",
        "chat":       "💬 محادثة",
    },
    "tr": {
        "image":      "🖼 Görsel oluşturma",
        "present":    "📊 Sunum/Slayt",
        "referat":    "📝 Rapor",
        "tezis":      "📄 Tez",
        "article":    "📰 Makale",
        "mustaqil":   "🗂 Bağımsız çalışma",
        "kurs":       "🎓 Dönem ödevi",
        "keys":       "💼 Vaka çalışması",
        "resume":     "📋 Özgeçmiş",
        "texmap":     "🗺 Teknoloji haritası",
        "quiz":       "❓ Quiz/Test",
        "crossword":  "🔠 Bulmaca",
        "translate":  "🌐 Çeviri",
        "essay":      "✍️ Deneme",
        "flashcard":  "🃏 Flaş kart",
        "convert":    "🔄 Dosya dönüştürme",
        "chat":       "💬 Sohbet",
    },
    "es": {
        "image":      "🖼 Generar imagen",
        "present":    "📊 Presentación/Diapositivas",
        "referat":    "📝 Informe",
        "tezis":      "📄 Tesis",
        "article":    "📰 Artículo",
        "mustaqil":   "🗂 Trabajo independiente",
        "kurs":       "🎓 Trabajo de curso",
        "keys":       "💼 Estudio de caso",
        "resume":     "📋 Currículum",
        "texmap":     "🗺 Mapa tecnológico",
        "quiz":       "❓ Quiz/Test",
        "crossword":  "🔠 Crucigrama",
        "translate":  "🌐 Traducción",
        "essay":      "✍️ Ensayo",
        "flashcard":  "🃏 Tarjeta flash",
        "convert":    "🔄 Convertir archivo",
        "chat":       "💬 Chat",
    },
    "hi": {
        "image":      "🖼 छवि बनाएं",
        "present":    "📊 प्रेजेंटेशन/स्लाइड",
        "referat":    "📝 रिपोर्ट",
        "tezis":      "📄 थीसिस",
        "article":    "📰 लेख",
        "mustaqil":   "🗂 स्वतंत्र कार्य",
        "kurs":       "🎓 कोर्स वर्क",
        "keys":       "💼 केस स्टडी",
        "resume":     "📋 रेज़्युमे",
        "texmap":     "🗺 तकनीकी मानचित्र",
        "quiz":       "❓ क्विज़/टेस्ट",
        "crossword":  "🔠 क्रॉसवर्ड",
        "translate":  "🌐 अनुवाद",
        "essay":      "✍️ निबंध",
        "flashcard":  "🃏 फ्लैश कार्ड",
        "convert":    "🔄 फ़ाइल कनवर्ट",
        "chat":       "💬 चैट",
    },
}

LIMIT_MSG = {
    "uz": (
        "⚠️ Kunlik limitingiz tugadi!\n\n"
        "🆓 Bepul: 1 ta so'rov/kun\n"
        "⭐ Standart: 20 ta so'rov/kun\n"
        "💎 Premium: Cheksiz\n\n"
        "Tarifni oshirish uchun 💎 Tariflar bo'limiga o'ting."
    ),
    "en": (
        "⚠️ Daily limit reached!\n\n"
        "🆓 Free: 1 request/day\n"
        "⭐ Standard: 20 requests/day\n"
        "💎 Premium: Unlimited\n\n"
        "Go to 💎 Tariffs to upgrade."
    ),
    "ru": (
        "⚠️ Дневной лимит исчерпан!\n\n"
        "🆓 Бесплатно: 1 запрос/день\n"
        "⭐ Стандарт: 20 запросов/день\n"
        "💎 Премиум: Безлимит\n\n"
        "Перейдите в 💎 Тарифы для повышения."
    ),
    "ar": (
        "⚠️ تم استنفاد الحد اليومي!\n\n"
        "🆓 مجاني: طلب واحد/يوم\n"
        "⭐ قياسي: 20 طلباً/يوم\n"
        "💎 بريميوم: غير محدود\n\n"
        "اذهب إلى 💎 الباقات للترقية."
    ),
    "tr": (
        "⚠️ Günlük limit doldu!\n\n"
        "🆓 Ücretsiz: 1 istek/gün\n"
        "⭐ Standart: 20 istek/gün\n"
        "💎 Premium: Sınırsız\n\n"
        "Yükseltmek için 💎 Tarifler'e gidin."
    ),
    "es": (
        "⚠️ ¡Límite diario alcanzado!\n\n"
        "🆓 Gratis: 1 solicitud/día\n"
        "⭐ Estándar: 20 solicitudes/día\n"
        "💎 Premium: Ilimitado\n\n"
        "Ve a 💎 Tarifas para mejorar."
    ),
    "hi": (
        "⚠️ दैनिक सीमा समाप्त!\n\n"
        "🆓 मुफ्त: 1 अनुरोध/दिन\n"
        "⭐ स्टैंडर्ड: 20 अनुरोध/दिन\n"
        "💎 प्रीमियम: असीमित\n\n"
        "अपग्रेड के लिए 💎 टैरिफ पर जाएं।"
    ),
}

PREMIUM_ONLY_MSG = {
    "uz": "💎 Bu funksiya faqat Premium tarif uchun!\n\n💎 Tariflar bo'limiga o'ting.",
    "en": "💎 This function is for Premium only!\n\nGo to 💎 Tariffs.",
    "ru": "💎 Эта функция только для Премиум!\n\nПерейдите в 💎 Тарифы.",
    "ar": "💎 هذه الوظيفة للبريميوم فقط!\n\nاذهب إلى 💎 الباقات.",
    "tr": "💎 Bu fonksiyon sadece Premium için!\n\n💎 Tariflere gidin.",
    "es": "💎 ¡Esta función es solo para Premium!\n\nVe a 💎 Tarifas.",
    "hi": "💎 यह फ़ंक्शन केवल प्रीमियम के लिए!\n\n💎 टैरिफ पर जाएं।",
}

THINKING_MSG = {
    "uz": "⏳ Tayyorlanmoqda...",
    "en": "⏳ Preparing...",
    "ru": "⏳ Подготовка...",
    "ar": "⏳ جارٍ التحضير...",
    "tr": "⏳ Hazırlanıyor...",
    "es": "⏳ Preparando...",
    "hi": "⏳ तैयारी हो रही है...",
}

BACK_BTN = {
    "uz": "🔙 Orqaga", "en": "🔙 Back", "ru": "🔙 Назад",
    "ar": "🔙 رجوع",  "tr": "🔙 Geri",  "es": "🔙 Atrás", "hi": "🔙 वापस"
}

PROMPT_MSG = {
    "image":     {"uz":"🖼 Rasm tavsifini yozing:","en":"🖼 Describe the image:","ru":"🖼 Опишите изображение:","ar":"🖼 صف الصورة:","tr":"🖼 Görseli tanımlayın:","es":"🖼 Describe la imagen:","hi":"🖼 छवि का वर्णन करें:"},
    "present":   {"uz":"📊 Taqdimot mavzusini yozing:","en":"📊 Write the presentation topic:","ru":"📊 Напишите тему презентации:","ar":"📊 اكتب موضوع العرض:","tr":"📊 Sunum konusunu yazın:","es":"📊 Escribe el tema de la presentación:","hi":"📊 प्रेजेंटेशन का विषय लिखें:"},
    "referat":   {"uz":"📝 Referat mavzusini yozing:","en":"📝 Write the report topic:","ru":"📝 Напишите тему реферата:","ar":"📝 اكتب موضوع التقرير:","tr":"📝 Rapor konusunu yazın:","es":"📝 Escribe el tema del informe:","hi":"📝 रिपोर्ट का विषय लिखें:"},
    "tezis":     {"uz":"📄 Tezis mavzusini yozing:","en":"📄 Write the thesis topic:","ru":"📄 Напишите тему тезиса:","ar":"📄 اكتب موضوع الأطروحة:","tr":"📄 Tez konusunu yazın:","es":"📄 Escribe el tema de la tesis:","hi":"📄 थीसिस का विषय लिखें:"},
    "article":   {"uz":"📰 Maqola mavzusini yozing:","en":"📰 Write the article topic:","ru":"📰 Напишите тему статьи:","ar":"📰 اكتب موضوع المقال:","tr":"📰 Makale konusunu yazın:","es":"📰 Escribe el tema del artículo:","hi":"📰 लेख का विषय लिखें:"},
    "mustaqil":  {"uz":"🗂 Mustaqil ish mavzusini yozing:","en":"🗂 Write the independent work topic:","ru":"🗂 Напишите тему самостоятельной работы:","ar":"🗂 اكتب موضوع العمل المستقل:","tr":"🗂 Bağımsız çalışma konusunu yazın:","es":"🗂 Escribe el tema del trabajo independiente:","hi":"🗂 स्वतंत्र कार्य का विषय लिखें:"},
    "kurs":      {"uz":"🎓 Kurs ishi mavzusini yozing:","en":"🎓 Write the course work topic:","ru":"🎓 Напишите тему курсовой:","ar":"🎓 اكتب موضوع العمل الدراسي:","tr":"🎓 Dönem ödevi konusunu yazın:","es":"🎓 Escribe el tema del trabajo de curso:","hi":"🎓 कोर्स वर्क का विषय लिखें:"},
    "keys":      {"uz":"💼 Keys mavzusini yozing:","en":"💼 Write the case study topic:","ru":"💼 Напишите тему кейса:","ar":"💼 اكتب موضوع دراسة الحالة:","tr":"💼 Vaka konusunu yazın:","es":"💼 Escribe el tema del caso:","hi":"💼 केस स्टडी का विषय लिखें:"},
    "resume":    {"uz":"📋 Kasb/soha va tajribangizni yozing:","en":"📋 Write your profession and experience:","ru":"📋 Напишите профессию и опыт:","ar":"📋 اكتب مهنتك وخبرتك:","tr":"📋 Meslek ve deneyiminizi yazın:","es":"📋 Escribe tu profesión y experiencia:","hi":"📋 अपना पेशा और अनुभव लिखें:"},
    "texmap":    {"uz":"🗺 Texnologik xarita mavzusini yozing:","en":"🗺 Write the tech map topic:","ru":"🗺 Напишите тему технологической карты:","ar":"🗺 اكتب موضوع الخريطة التقنية:","tr":"🗺 Teknoloji haritası konusunu yazın:","es":"🗺 Escribe el tema del mapa tecnológico:","hi":"🗺 तकनीकी मानचित्र का विषय लिखें:"},
    "quiz":      {"uz":"❓ Quiz mavzusini va savollar sonini yozing:\nMisol: Matematika 10 ta savol","en":"❓ Write quiz topic and number of questions:\nExample: Math 10 questions","ru":"❓ Напишите тему и количество вопросов:\nПример: Математика 10 вопросов","ar":"❓ اكتب موضوع الاختبار وعدد الأسئلة:\nمثال: رياضيات 10 أسئلة","tr":"❓ Quiz konusunu ve soru sayısını yazın:\nÖrnek: Matematik 10 soru","es":"❓ Escribe el tema del quiz y número de preguntas:\nEjemplo: Matemáticas 10 preguntas","hi":"❓ क्विज़ विषय और प्रश्नों की संख्या लिखें:\nउदाहरण: गणित 10 प्रश्न"},
    "crossword": {"uz":"🔠 Krossvord mavzusini yozing:","en":"🔠 Write the crossword topic:","ru":"🔠 Напишите тему кроссворда:","ar":"🔠 اكتب موضوع الكلمات المتقاطعة:","tr":"🔠 Bulmaca konusunu yazın:","es":"🔠 Escribe el tema del crucigrama:","hi":"🔠 क्रॉसवर्ड का विषय लिखें:"},
    "translate": {"uz":"🌐 Tarjima qilinadigan matnni yozing va tilni ko'rsating:\nMisol: Hello world — O'zbekchaga","en":"🌐 Write the text to translate and target language:\nExample: Salom dunyo — to English","ru":"🌐 Напишите текст и язык перевода:\nПример: Hello world — на русский","ar":"🌐 اكتب النص واللغة المستهدفة:\nمثال: مرحبا — إلى الإنجليزية","tr":"🌐 Çevrilecek metni ve hedef dili yazın:\nÖrnek: Merhaba dünya — İngilizceye","es":"🌐 Escribe el texto a traducir y el idioma:\nEjemplo: Hola mundo — al inglés","hi":"🌐 अनुवाद करने के लिए पाठ और भाषा लिखें:\nउदाहरण: Hello world — हिंदी में"},
    "essay":     {"uz":"✍️ Insho/Esse mavzusini yozing:","en":"✍️ Write the essay topic:","ru":"✍️ Напишите тему эссе:","ar":"✍️ اكتب موضوع المقال الإنشائي:","tr":"✍️ Deneme konusunu yazın:","es":"✍️ Escribe el tema del ensayo:","hi":"✍️ निबंध का विषय लिखें:"},
    "flashcard": {"uz":"🃏 Flesh karta mavzusini yozing:","en":"🃏 Write the flashcard topic:","ru":"🃏 Напишите тему флэш-карты:","ar":"🃏 اكتب موضوع البطاقة التعليمية:","tr":"🃏 Flaş kart konusunu yazın:","es":"🃏 Escribe el tema de la tarjeta flash:","hi":"🃏 फ्लैश कार्ड का विषय लिखें:"},
    "chat":      {"uz":"💬 Savolingizni yozing:","en":"💬 Write your question:","ru":"💬 Напишите вопрос:","ar":"💬 اكتب سؤالك:","tr":"💬 Sorunuzu yazın:","es":"💬 Escribe tu pregunta:","hi":"💬 अपना प्रश्न लिखें:"},
    "convert":   {"uz":"🔄 Faylni yuboring (PDF, Word, Excel, Rasm...):","en":"🔄 Send the file (PDF, Word, Excel, Image...):","ru":"🔄 Отправьте файл (PDF, Word, Excel, Фото...):","ar":"🔄 أرسل الملف (PDF، Word، Excel، صورة...):","tr":"🔄 Dosyayı gönderin (PDF, Word, Excel, Resim...):","es":"🔄 Envía el archivo (PDF, Word, Excel, Imagen...):","hi":"🔄 फ़ाइल भेजें (PDF, Word, Excel, छवि...):"},
}

# ============================================================
# FSM
# ============================================================
class AIS(StatesGroup):
    waiting_input  = State()
    waiting_file   = State()
    chat           = State()

# ============================================================
# AI MENU KLAVIATURASI
# ============================================================
def ai_menu_kb(lang):
    import sys
    main_module = sys.modules.get('__main__')
    sett   = main_module.load_settings() if main_module and hasattr(main_module, 'load_settings') else {}
    custom = sett.get("custom_section_texts", {})

    funcs  = AI_FUNCTIONS.get(lang, AI_FUNCTIONS["en"])
    btn_map = {
        "image": "btn_ai_image", "present": "btn_ai_present",
        "referat": "btn_ai_referat", "tezis": "btn_ai_tezis",
        "article": "btn_ai_article", "mustaqil": "btn_ai_mustaqil",
        "kurs": "btn_ai_kurs", "keys": "btn_ai_keys",
        "resume": "btn_ai_resume", "texmap": "btn_ai_texmap",
        "quiz": "btn_ai_quiz", "crossword": "btn_ai_crossword",
        "translate": "btn_ai_translate", "essay": "btn_ai_essay",
        "flashcard": "btn_ai_flashcard", "convert": "btn_ai_convert",
        "chat": "btn_ai_chat",
    }
    btns = []
    keys = list(funcs.keys())
    for i in range(0, len(keys) - 1, 2):
        k1 = keys[i]
        k2 = keys[i+1]
        t1 = custom.get(btn_map.get(k1,""), {}).get(lang, funcs[k1])
        t2 = custom.get(btn_map.get(k2,""), {}).get(lang, funcs[k2])
        btns.append([
            InlineKeyboardButton(text=t1, callback_data=f"ai_{k1}"),
            InlineKeyboardButton(text=t2, callback_data=f"ai_{k2}"),
        ])
    if len(keys) % 2:
        k = keys[-1]
        t = custom.get(btn_map.get(k,""), {}).get(lang, funcs[k])
        btns.append([InlineKeyboardButton(text=t, callback_data=f"ai_{k}")])
    btns.append([InlineKeyboardButton(
        text=BACK_BTN.get(lang, "🔙 Back"),
        callback_data="ai_back"
    )])
    return InlineKeyboardMarkup(inline_keyboard=btns)

def back_to_ai_kb(lang):
    """AI funksiyasi ichidan AI menyusiga qaytish tugmasi."""
    return InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(
            text=BACK_BTN.get(lang, "🔙 Back"),
            callback_data="ai_menu"
        )
    ]])

# ============================================================
# AI MODEL TANLASH
# ============================================================
async def get_ai_response(prompt, system, tariff, lang):
    if tariff == "premium":
        # GPT-4o
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    OPENAI_URL,
                    headers={
                        "Authorization": f"Bearer {OPENAI_API}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "gpt-4o",
                        "messages": [
                            {"role": "system", "content": system},
                            {"role": "user",   "content": prompt}
                        ],
                        "max_tokens": 4096,
                        "temperature": 0.7
                    }
                ) as r:
                    data = await r.json()
            return data["choices"][0]["message"]["content"]
        except:
            pass

    elif tariff == "standard":
        # GPT-4o mini
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    OPENAI_URL,
                    headers={
                        "Authorization": f"Bearer {OPENAI_API}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "gpt-4o-mini",
                        "messages": [
                            {"role": "system", "content": system},
                            {"role": "user",   "content": prompt}
                        ],
                        "max_tokens": 2048,
                        "temperature": 0.7
                    }
                ) as r:
                    data = await r.json()
            return data["choices"][0]["message"]["content"]
        except:
            pass

    # Groq (bepul tarif yoki fallback)
    async with aiohttp.ClientSession() as session:
        async with session.post(
            GROQ_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": GROQ_MODEL,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user",   "content": prompt}
                ],
                "max_tokens": 2048,
                "temperature": 0.7
            }
        ) as r:
            data = await r.json()
    return data["choices"][0]["message"]["content"]


# ============================================================
# RASM GENERATSIYA (DALL-E 3)
# ============================================================
async def generate_image(prompt):
    async with aiohttp.ClientSession() as session:
        async with session.post(
            DALLE_URL,
            headers={
                "Authorization": f"Bearer {OPENAI_API}",
                "Content-Type": "application/json"
            },
            json={
                "model": "dall-e-3",
                "prompt": prompt,
                "n": 1,
                "size": "1024x1024",
                "quality": "standard"
            }
        ) as r:
            data = await r.json()
    return data["data"][0]["url"]


# ============================================================
# SYSTEM PROMPTLAR
# ============================================================
def get_system_prompt(func, lang):
    lang_name = {
        "uz": "o'zbek", "en": "English", "ru": "русском",
        "ar": "العربية", "tr": "Türkçe", "es": "español", "hi": "हिंदी"
    }.get(lang, "English")

    prompts = {
        "referat": f"Siz professional referat yozuvchisiz. Foydalanuvchi bergan mavzu bo'yicha to'liq, ilmiy, tarkibli referat yozing. Til: {lang_name}. Tarkib: Kirish, Asosiy qism (3-4 bo'lim), Xulosa, Adabiyotlar ro'yxati.",
        "tezis": f"Siz ilmiy tezis yozuvchisiz. Mavzu bo'yicha qisqa, aniq, ilmiy tezis yozing. Til: {lang_name}. Tarkib: Muammo, Maqsad, Vazifalar, Metodlar, Kutilayotgan natijalar.",
        "article": f"Siz professional maqola yozuvchisiz. Mavzu bo'yicha to'liq, mazmunli maqola yozing. Til: {lang_name}. Tarkib: Sarlavha, Kirish, Asosiy qism, Xulosa.",
        "mustaqil": f"Siz mustaqil ish yozuvchisiz. Mavzu bo'yicha to'liq mustaqil ish yozing. Til: {lang_name}. Tarkib: Kirish, Nazariy qism, Amaliy qism, Xulosa.",
        "kurs": f"Siz kurs ishi yozuvchisiz. Mavzu bo'yicha to'liq kurs ishi yozing. Til: {lang_name}. Tarkib: Mundarija, Kirish, I bob, II bob, Xulosa, Adabiyotlar.",
        "keys": f"Siz biznes keys (case study) yozuvchisiz. Mavzu bo'yicha to'liq keys yozing. Til: {lang_name}. Tarkib: Muammo tavsifi, Tahlil, Yechim variantlari, Tavsiya.",
        "resume": f"Siz professional rezyume yozuvchisiz. Foydalanuvchi ma'lumotlari asosida professional rezyume tuzing. Til: {lang_name}. Tarkib: Shaxsiy ma'lumot, Tajriba, Ta'lim, Ko'nikmalar, Qo'shimcha.",
        "texmap": f"Siz texnologik xarita tuzuvchisiz. Mavzu bo'yicha to'liq texnologik xarita tuzing. Til: {lang_name}. Tarkib: Dars maqsadi, Bosqichlar, Vaqt, Metodlar, Natija.",
        "quiz": f"Siz test/quiz tuzuvchisiz. Mavzu va son bo'yicha test tuzing. Har bir savolda 4 ta javob varianti bo'lsin, to'g'ri javobni belgilang. Til: {lang_name}.",
        "crossword": f"Siz krossvord tuzuvchisiz. Mavzu bo'yicha krossvord tuzing. Savol va javoblarni aniq ko'rsating. Til: {lang_name}.",
        "translate": f"Siz professional tarjimonisiz. Foydalanuvchi bergan matnni ko'rsatilgan tilga to'liq va aniq tarjima qiling.",
        "essay": f"Siz insho/esse yozuvchisiz. Mavzu bo'yicha to'liq, mazmunli insho yozing. Til: {lang_name}. Tarkib: Kirish, Asosiy fikr, Dalillar, Xulosa.",
        "flashcard": f"Siz flesh karta tuzuvchisiz. Mavzu bo'yicha 10-15 ta flesh karta tuzing. Har bir kartada: SAVOL va JAVOB bo'lsin. Til: {lang_name}.",
        "present": f"Siz taqdimot tuzuvchisiz. Mavzu bo'yicha 10-15 ta slayd uchun kontent tuzing. Har bir slaydda: Sarlavha va asosiy fikrlar. Til: {lang_name}.",
        "chat": f"Siz yordamchi assistantsiz. Foydalanuvchi savollariga aniq, foydali javob bering. Til: {lang_name}.",
    }
    return prompts.get(func, prompts["chat"])


# ============================================================
# WORD FAYL YARATISH
# ============================================================
async def create_word_file(content, filename):
    try:
        from docx import Document
        doc = Document()
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("• "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        path = f"/tmp/{filename}.docx"
        doc.save(path)
        return path
    except Exception as e:
        return None


# ============================================================
# FAYL KONVERTATSIYA
# ============================================================
CONVERT_FORMATS = {
    "uz": {
        "pdf_word":   "PDF → Word",
        "word_pdf":   "Word → PDF",
        "img_pdf":    "Rasm → PDF",
        "pdf_img":    "PDF → Rasm",
        "pdf_txt":    "PDF → Matn",
        "excel_pdf":  "Excel → PDF",
        "pdf_excel":  "PDF → Excel",
        "ppt_pdf":    "PowerPoint → PDF",
        "png_jpg":    "PNG → JPG",
        "jpg_png":    "JPG → PNG",
        "mp4_mp3":    "Video → Audio",
        "pdf_speech": "PDF → Matn (OCR)",
    },
    "en": {
        "pdf_word":   "PDF → Word",
        "word_pdf":   "Word → PDF",
        "img_pdf":    "Image → PDF",
        "pdf_img":    "PDF → Image",
        "pdf_txt":    "PDF → Text",
        "excel_pdf":  "Excel → PDF",
        "pdf_excel":  "PDF → Excel",
        "ppt_pdf":    "PowerPoint → PDF",
        "png_jpg":    "PNG → JPG",
        "jpg_png":    "JPG → PNG",
        "mp4_mp3":    "Video → Audio",
        "pdf_speech": "PDF → Text (OCR)",
    },
    "ru": {
        "pdf_word":   "PDF → Word",
        "word_pdf":   "Word → PDF",
        "img_pdf":    "Фото → PDF",
        "pdf_img":    "PDF → Фото",
        "pdf_txt":    "PDF → Текст",
        "excel_pdf":  "Excel → PDF",
        "pdf_excel":  "PDF → Excel",
        "ppt_pdf":    "PowerPoint → PDF",
        "png_jpg":    "PNG → JPG",
        "jpg_png":    "JPG → PNG",
        "mp4_mp3":    "Видео → Аудио",
        "pdf_speech": "PDF → Текст (OCR)",
    },
    "ar": {
        "pdf_word":   "PDF → Word",
        "word_pdf":   "Word → PDF",
        "img_pdf":    "صورة → PDF",
        "pdf_img":    "PDF → صورة",
        "pdf_txt":    "PDF → نص",
        "excel_pdf":  "Excel → PDF",
        "pdf_excel":  "PDF → Excel",
        "ppt_pdf":    "PowerPoint → PDF",
        "png_jpg":    "PNG → JPG",
        "jpg_png":    "JPG → PNG",
        "mp4_mp3":    "فيديو → صوت",
        "pdf_speech": "PDF → نص (OCR)",
    },
    "tr": {
        "pdf_word":   "PDF → Word",
        "word_pdf":   "Word → PDF",
        "img_pdf":    "Resim → PDF",
        "pdf_img":    "PDF → Resim",
        "pdf_txt":    "PDF → Metin",
        "excel_pdf":  "Excel → PDF",
        "pdf_excel":  "PDF → Excel",
        "ppt_pdf":    "PowerPoint → PDF",
        "png_jpg":    "PNG → JPG",
        "jpg_png":    "JPG → PNG",
        "mp4_mp3":    "Video → Ses",
        "pdf_speech": "PDF → Metin (OCR)",
    },
    "es": {
        "pdf_word":   "PDF → Word",
        "word_pdf":   "Word → PDF",
        "img_pdf":    "Imagen → PDF",
        "pdf_img":    "PDF → Imagen",
        "pdf_txt":    "PDF → Texto",
        "excel_pdf":  "Excel → PDF",
        "pdf_excel":  "PDF → Excel",
        "ppt_pdf":    "PowerPoint → PDF",
        "png_jpg":    "PNG → JPG",
        "jpg_png":    "JPG → PNG",
        "mp4_mp3":    "Video → Audio",
        "pdf_speech": "PDF → Texto (OCR)",
    },
    "hi": {
        "pdf_word":   "PDF → Word",
        "word_pdf":   "Word → PDF",
        "img_pdf":    "छवि → PDF",
        "pdf_img":    "PDF → छवि",
        "pdf_txt":    "PDF → पाठ",
        "excel_pdf":  "Excel → PDF",
        "pdf_excel":  "PDF → Excel",
        "ppt_pdf":    "PowerPoint → PDF",
        "png_jpg":    "PNG → JPG",
        "jpg_png":    "JPG → PNG",
        "mp4_mp3":    "वीडियो → ऑडियो",
        "pdf_speech": "PDF → पाठ (OCR)",
    },
}

def convert_formats_kb(lang):
    formats = CONVERT_FORMATS.get(lang, CONVERT_FORMATS["en"])
    btns    = []
    keys    = list(formats.keys())
    for i in range(0, len(keys) - 1, 2):
        btns.append([
            InlineKeyboardButton(text=formats[keys[i]],   callback_data=f"conv_fmt_{keys[i]}"),
            InlineKeyboardButton(text=formats[keys[i+1]], callback_data=f"conv_fmt_{keys[i+1]}"),
        ])
    if len(keys) % 2:
        btns.append([InlineKeyboardButton(
            text=formats[keys[-1]],
            callback_data=f"conv_fmt_{keys[-1]}"
        )])
    btns.append([InlineKeyboardButton(
        text=BACK_BTN.get(lang, "🔙 Back"),
        callback_data="ai_menu"
    )])
    return InlineKeyboardMarkup(inline_keyboard=btns)


async def convert_file(message, fmt, bot):
    try:
        from PIL import Image
        import fitz  # PyMuPDF
        import io

        if fmt == "img_pdf":
            if not message.photo and not message.document:
                return None
            if message.photo:
                file = await bot.get_file(message.photo[-1].file_id)
            else:
                file = await bot.get_file(message.document.file_id)
            data = await bot.download_file(file.file_path)
            img  = Image.open(io.BytesIO(data.read())).convert("RGB")
            path = "/tmp/converted.pdf"
            img.save(path, "PDF")
            return path

        elif fmt == "pdf_img":
            if not message.document:
                return None
            file = await bot.get_file(message.document.file_id)
            data = await bot.download_file(file.file_path)
            doc  = fitz.open(stream=data.read(), filetype="pdf")
            page = doc[0]
            pix  = page.get_pixmap()
            path = "/tmp/converted.png"
            pix.save(path)
            return path

        elif fmt == "pdf_txt":
            if not message.document:
                return None
            file = await bot.get_file(message.document.file_id)
            data = await bot.download_file(file.file_path)
            doc  = fitz.open(stream=data.read(), filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            path = "/tmp/converted.txt"
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
            return path

        elif fmt == "png_jpg":
            if not message.photo and not message.document:
                return None
            if message.photo:
                file = await bot.get_file(message.photo[-1].file_id)
            else:
                file = await bot.get_file(message.document.file_id)
            data = await bot.download_file(file.file_path)
            img  = Image.open(io.BytesIO(data.read())).convert("RGB")
            path = "/tmp/converted.jpg"
            img.save(path, "JPEG")
            return path

        elif fmt == "jpg_png":
            if not message.photo and not message.document:
                return None
            if message.photo:
                file = await bot.get_file(message.photo[-1].file_id)
            else:
                file = await bot.get_file(message.document.file_id)
            data = await bot.download_file(file.file_path)
            img  = Image.open(io.BytesIO(data.read()))
            path = "/tmp/converted.png"
            img.save(path, "PNG")
            return path

        elif fmt == "word_pdf":
            if not message.document:
                return None
            file = await bot.get_file(message.document.file_id)
            data = await bot.download_file(file.file_path)
            inp  = "/tmp/input.docx"
            with open(inp, "wb") as f:
                f.write(data.read())
            out = "/tmp/converted.pdf"
            os.system(f"libreoffice --headless --convert-to pdf {inp} --outdir /tmp/")
            return out if os.path.exists(out) else None

        elif fmt == "pdf_word":
            if not message.document:
                return None
            file = await bot.get_file(message.document.file_id)
            data = await bot.download_file(file.file_path)
            doc  = fitz.open(stream=data.read(), filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            from docx import Document as DocxDoc
            d    = DocxDoc()
            for line in text.split("\n"):
                d.add_paragraph(line)
            path = "/tmp/converted.docx"
            d.save(path)
            return path

        elif fmt == "mp4_mp3":
            if not message.video and not message.document:
                return None
            if message.video:
                file = await bot.get_file(message.video.file_id)
            else:
                file = await bot.get_file(message.document.file_id)
            data = await bot.download_file(file.file_path)
            inp  = "/tmp/input.mp4"
            out  = "/tmp/converted.mp3"
            with open(inp, "wb") as f:
                f.write(data.read())
            os.system(f"ffmpeg -i {inp} -q:a 0 -map a {out} -y")
            return out if os.path.exists(out) else None

    except Exception as e:
     return None
# ============================================================
# HANDLERLARNI RO'YXATDAN O'TKAZISH
# ============================================================
def register_ai_handlers(dp, bot):

        # ---- AI MENYU ----
        @dp.callback_query(F.data == "ai_menu")
        async def cb_ai_menu(callback: types.CallbackQuery, state: FSMContext):
            uid = callback.from_user.id
            user = get_user(uid)
            lang = user.get("lang", "en")
            await state.clear()
            import sys
            main_module = sys.modules.get('__main__')
            sett = main_module.load_settings() if main_module and hasattr(main_module, 'load_settings') else {}
            custom_title = sett.get("custom_section_texts", {}).get("ai_menu_title", {}).get(lang, "")
            title = custom_title if custom_title else AI_MENU_TITLE.get(lang, AI_MENU_TITLE["en"])
            await callback.message.edit_text(
                title,
                reply_markup=ai_menu_kb(lang)
            )
        @dp.callback_query(F.data == "ai_back")
        async def cb_ai_back(callback: types.CallbackQuery, state: FSMContext):
            await state.clear()
            await callback.message.delete()
        # ---- AI FUNKSIYA TANLASH ----
        @dp.callback_query(F.data.startswith("ai_"))
        async def cb_ai_func(callback: types.CallbackQuery, state: FSMContext):
            uid = callback.from_user.id
            user = get_user(uid)
            lang = user.get("lang", "en")
            func = callback.data[3:]

            if func in ("menu", "back"):
                return

            tariff = check_subscription(uid)

            def pay_required_msg():
                bal = get_balance(uid)
                import sys
                main_module = sys.modules.get('__main__')
                sett = main_module.load_settings() if main_module and hasattr(main_module, 'load_settings') else {}
                custom_pay = sett.get("custom_section_texts", {}).get("ai_pay_msg", {}).get(lang, "")
                default_pay = {
                    "uz": f"💳 Bu funksiyadan bepul foydalandingiz.\n\nBalansingiz: {bal:,} so'm\nKeyingi foydalanish: 3,000 so'm\n\n💰 Balansni to'ldiring yoki 💎 tarif sotib oling.",
                    "en": f"💳 You've used the free trial.\n\nBalance: {bal:,} UZS\nNext use: 3,000 UZS\n\n💰 Top up balance or 💎 buy subscription.",
                    "ru": f"💳 Вы использовали бесплатный период.\n\nБаланс: {bal:,} UZS\nСледующее: 3,000 UZS\n\n💰 Пополните баланс или 💎 купите подписку.",
                    "ar": f"💳 استخدمت الفترة المجانية.\n\nالرصيد: {bal:,} UZS\nالتالي: 3,000 UZS\n\n💰 اشحن الرصيد أو 💎 اشترِ اشتراكاً.",
                    "tr": f"💳 Ücretsiz denemeyi kullandınız.\n\nBakiye: {bal:,} UZS\nSonraki: 3,000 UZS\n\n💰 Bakiye yükleyin veya 💎 abonelik alın.",
                    "es": f"💳 Usaste la prueba gratuita.\n\nSaldo: {bal:,} UZS\nSiguiente: 3,000 UZS\n\n💰 Recarga o 💎 compra suscripción.",
                    "hi": f"💳 मुफ्त परीक्षण उपयोग किया।\n\nबैलेंस: {bal:,} UZS\nअगला: 3,000 UZS\n\n💰 बैलेंस भरें या 💎 सदस्यता खरीदें।",
                }
                return custom_pay.replace("{bal}", f"{bal:,}") if custom_pay else default_pay.get(lang, default_pay["en"])

            # ---- Rasm yaratish: faqat Premium ----
            if func == "image":
                if tariff != "premium":
                    await callback.answer(PREMIUM_ONLY_MSG.get(lang, PREMIUM_ONLY_MSG["en"]), show_alert=True)
                    return
                allowed, _ = can_use_ai(uid, "image")
                if not allowed:
                    await callback.answer(LIMIT_MSG.get(lang, LIMIT_MSG["en"]), show_alert=True)
                    return
                await state.set_state(AIS.waiting_input)
                await state.update_data(func="image", lang=lang, tariff=tariff)
                await callback.message.edit_text(
                    PROMPT_MSG["image"].get(lang, PROMPT_MSG["image"]["en"]),
                    reply_markup=back_to_ai_kb(lang)
                )
                return

            # ---- Fayl konvertatsiya ----
            if func == "convert":
                await state.set_state(AIS.waiting_file)
                await state.update_data(func=func, lang=lang)
                await callback.message.edit_text(
                    CONVERT_FORMATS.get(lang, CONVERT_FORMATS["en"]).get("pdf_word", ""),
                    reply_markup=convert_formats_kb(lang)
                )
                return

            # ---- Boshqa funksiyalar: tarifga qarab cheklov ----
            if tariff == "premium":
                pass  # cheksiz
            elif tariff == "standard":
                allowed, _ = can_use_ai(uid, func)
                if not allowed:
                    await callback.answer(LIMIT_MSG.get(lang, LIMIT_MSG["en"]), show_alert=True)
                    return
            else:  # free
                if not has_free_usage(uid, func):
                    mark_free_usage(uid, func)
                elif has_enough_balance(uid, 3000):
                    deduct_balance(uid, 3000)
                else:
                    await callback.answer(pay_required_msg(), show_alert=True)
                    return

            await state.set_state(AIS.waiting_input)
            await state.update_data(func=func, lang=lang, tariff=tariff)
            import sys
            main_module = sys.modules.get('__main__')
            sett = main_module.load_settings() if main_module and hasattr(main_module, 'load_settings') else {}
            custom_prompt = sett.get("custom_section_texts", {}).get(f"ai_prompt_{func}", {}).get(lang, "")
            prompt_text = custom_prompt if custom_prompt else PROMPT_MSG.get(func, {}).get(
                lang, PROMPT_MSG.get(func, {}).get("en", ""))
            await callback.message.edit_text(
                prompt_text,
                reply_markup=back_to_ai_kb(lang)
            )

        # ---- KONVERT FORMAT TANLASH ----
        @dp.callback_query(F.data.startswith("conv_fmt_"))
        async def cb_conv_fmt(callback: types.CallbackQuery, state: FSMContext):
            uid = callback.from_user.id
            user = get_user(uid)
            lang = user.get("lang", "en")
            fmt = callback.data[9:]
            await state.set_state(AIS.waiting_file)
            await state.update_data(func="convert", fmt=fmt, lang=lang)
            send_file_msg = PROMPT_MSG.get("convert", {}).get(lang, PROMPT_MSG["convert"]["en"])
            await callback.message.edit_text(
                send_file_msg,
                reply_markup=back_to_ai_kb(lang)
            )
        # ---- FAYL QABUL QILISH (KONVERT) ----
        @dp.message(AIS.waiting_file)
        async def handle_file_convert(message: types.Message, state: FSMContext):
            uid = message.from_user.id
            user = get_user(uid)
            data = await state.get_data()
            lang = data.get("lang", "en")
            fmt = data.get("fmt", "")

            if not any([
                message.document, message.photo,
                message.video, message.audio
            ]):
                await message.answer(PROMPT_MSG["convert"].get(lang, ""))
                return

            thinking = await message.answer(THINKING_MSG.get(lang, "⏳"))

            result_path = await convert_file(message, fmt, bot)

            await thinking.delete()

            if not result_path or not os.path.exists(result_path):
                err = {
                    "uz": "❌ Konvertatsiya xatoligi. Boshqa fayl yuboring.",
                    "en": "❌ Conversion error. Please send another file.",
                    "ru": "❌ Ошибка конвертации. Отправьте другой файл.",
                    "ar": "❌ خطأ في التحويل. أرسل ملفاً آخر.",
                    "tr": "❌ Dönüştürme hatası. Başka dosya gönderin.",
                    "es": "❌ Error de conversión. Envía otro archivo.",
                    "hi": "❌ रूपांतरण त्रुटि। दूसरी फ़ाइल भेजें।",
                }
                await message.answer(err.get(lang, err["en"]))
                return

            log_ai_request(uid)
            await state.clear()

            done = {
                "uz": "✅ Tayyor!",
                "en": "✅ Done!",
                "ru": "✅ Готово!",
                "ar": "✅ تم!",
                "tr": "✅ Hazır!",
                "es": "✅ ¡Listo!",
                "hi": "✅ तैयार!",
            }
            with open(result_path, "rb") as f:
                await message.answer_document(f, caption=done.get(lang, "✅"))

        # ---- AI INPUT QABUL QILISH ----
        @dp.message(AIS.waiting_input)
        async def handle_ai_input(message: types.Message, state: FSMContext):
            uid = message.from_user.id
            user = get_user(uid)
            data = await state.get_data()
            func = data.get("func", "chat")
            lang = data.get("lang", "en")
            tariff = data.get("tariff", "free")
            text = message.text or ""

            if not text:
                await message.answer(PROMPT_MSG.get(func, {}).get(lang, ""))
                return

            # ---- Rasm yaratish (DALL-E) ----
            if func == "image":
                thinking = await message.answer(THINKING_MSG.get(lang, "⏳"))
                try:
                    url = await generate_image(text)
                    await thinking.delete()
                    log_ai_request(uid)
                    await state.clear()
                    done = {
                        "uz": "🖼 Rasm tayyor!", "en": "🖼 Image ready!",
                        "ru": "🖼 Изображение готово!", "ar": "🖼 الصورة جاهزة!",
                        "tr": "🖼 Görsel hazır!", "es": "🖼 ¡Imagen lista!",
                        "hi": "🖼 छवि तैयार!",
                    }
                    await message.answer_photo(url, caption=done.get(lang, done["en"]))
                    await message.answer(done.get(lang, done["en"]), reply_markup=back_to_ai_kb(lang))
                except Exception:
                    await thinking.delete()
                    err = {
                        "uz": "❌ Rasm yaratishda xatolik.", "en": "❌ Image generation error.",
                        "ru": "❌ Ошибка генерации изображения.", "ar": "❌ خطأ في توليد الصورة.",
                        "tr": "❌ Görsel oluşturma hatası.", "es": "❌ Error al generar imagen.",
                        "hi": "❌ छवि निर्माण त्रुटि।",
                    }
                    await message.answer(err.get(lang, err["en"]), reply_markup=back_to_ai_kb(lang))
                return

            thinking = await message.answer(THINKING_MSG.get(lang, "⏳"))

            try:
                system = get_system_prompt(func, lang)
                response = await get_ai_response(text, system, tariff, lang)

                await thinking.delete()
                log_ai_request(uid)

                # Chat uchun holatni saqlaymiz
                if func == "chat":
                    await message.answer(
                        f"🤖 {response}",
                        reply_markup=back_to_ai_kb(lang)
                    )
                    return

                # Boshqa funksiyalar uchun Word ham yuboramiz
                await message.answer(
                    f"✅\n\n{response[:3000]}",
                    reply_markup=back_to_ai_kb(lang)
                )

                # Word fayl yaratish
                func_name = AI_FUNCTIONS.get(lang, AI_FUNCTIONS["en"]).get(func, func)
                word_path = await create_word_file(response, f"{func}_{uid}")
                if word_path and os.path.exists(word_path):
                    word_caption = {
                        "uz": f"📄 {func_name} — Word fayl",
                        "en": f"📄 {func_name} — Word file",
                        "ru": f"📄 {func_name} — Word файл",
                        "ar": f"📄 {func_name} — ملف Word",
                        "tr": f"📄 {func_name} — Word dosyası",
                        "es": f"📄 {func_name} — Archivo Word",
                        "hi": f"📄 {func_name} — Word फ़ाइल",
                    }
                    with open(word_path, "rb") as f:
                        await message.answer_document(
                            f,
                            caption=word_caption.get(lang, word_caption["en"])
                        )

                await state.clear()

            except Exception as e:
                await thinking.delete()
                err = {
                    "uz": "❌ Xatolik yuz berdi. Qaytadan urinib ko'ring.",
                    "en": "❌ An error occurred. Please try again.",
                    "ru": "❌ Произошла ошибка. Попробуйте снова.",
                    "ar": "❌ حدث خطأ. حاول مرة أخرى.",
                    "tr": "❌ Bir hata oluştu. Tekrar deneyin.",
                    "es": "❌ Ocurrió un error. Inténtalo de nuevo.",
                    "hi": "❌ एक त्रुटि हुई। फिर से प्रयास करें।",
                }
                await message.answer(err.get(lang, err["en"]))
